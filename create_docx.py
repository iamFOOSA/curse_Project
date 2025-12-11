#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_code_block(doc, code, filename=""):
    if filename:
        p = doc.add_paragraph()
        run = p.add_run(f"Файл {filename}:")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()

def read_file_content(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Ошибка чтения файла: {e}"

def main():
    doc = Document()
    
    title = doc.add_heading('Листинг кода всех функций (кроме Qt)', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    sources_dir = os.path.join(base_dir, "sources")
    
    headers_dir = os.path.join(base_dir, "headers")
    
    cpp_files = [
        "user.cpp",
        "food.cpp",
        "manager.cpp",
        "recommendations.cpp",
        "nutrition_manager.cpp",
        "nutrition_advisor.cpp",
        "meal_tracker.cpp",
    ]
    
    h_files = [
        "user.h",
        "food.h",
        "manager.h",
        "recommendations.h",
        "nutrition_manager.h",
        "nutrition_advisor.h",
        "meal_tracker.h",
    ]
    
    doc.add_heading('Заголовочные файлы (.h)', 1)
    doc.add_paragraph()
    
    for filename in h_files:
        filepath = os.path.join(headers_dir, filename)
        if os.path.exists(filepath):
            code = read_file_content(filepath)
            add_code_block(doc, code, f"headers/{filename}")
        else:
            print(f"Файл не найден: {filepath}")
    
    add_page_break(doc)
    doc.add_heading('Исходные файлы (.cpp)', 1)
    doc.add_paragraph()
    
    for filename in cpp_files:
        filepath = os.path.join(sources_dir, filename)
        if os.path.exists(filepath):
            code = read_file_content(filepath)
            add_code_block(doc, code, f"sources/{filename}")
        else:
            print(f"Файл не найден: {filepath}")
    
    output_file = os.path.join(base_dir, "code_listing.docx")
    doc.save(output_file)
    print(f"Документ создан: {output_file}")

if __name__ == "__main__":
    main()
